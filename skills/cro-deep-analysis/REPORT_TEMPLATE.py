#!/usr/bin/env python3
"""CROディープ分析レポート — Word (.docx) + グラフ生成テンプレート

使い方:
1. DATA dict に Phase 2 の分析結果を注入
2. NARRATIVES dict に各セクションの分析文章を注入
3. python3 で実行 → .docx が出力される

依存パッケージ（venv推奨）:
  pip install python-docx matplotlib numpy
"""

import os
import sys
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.patches as mpatches
import numpy as np
from matplotlib.patches import Polygon
from matplotlib.colors import LinearSegmentedColormap

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ================================================================
# COLOR PALETTE
# ================================================================
C = {
    "primary": "#0f3460", "accent": "#e94560", "green": "#27ae60",
    "yellow": "#f39c12", "red": "#e74c3c", "gray": "#95a5a6",
    "dark": "#2c3e50", "light": "#ecf0f1",
    "green_bg": "#d5f5e3", "red_bg": "#fdedec",
    "yellow_bg": "#fef9e7", "blue_bg": "#d6eaf8",
}

# ================================================================
# matplotlib SETTINGS (per CHART_SPEC.md)
# ================================================================
plt.rcParams.update({
    "font.family": "Hiragino Sans",
    "font.size": 11,
    "axes.titlesize": 14,
    "axes.titleweight": "bold",
    "figure.facecolor": "#ffffff",
    "axes.facecolor": "#fafafa",
    "axes.edgecolor": "#dddddd",
    "axes.grid": True,
    "grid.alpha": 0.3,
    "grid.color": "#cccccc",
})


# ================================================================
# DATA INJECTION POINT
# Phase 2 の分析結果をここに注入する
# ================================================================
DATA = {
    # -- 基本情報 --
    "site_name": "SITE_NAME",
    "site_url": "www.example.com",
    "ga4_property": "000000000",
    "period": "過去30日間（クリーンデータ適用済み）",
    "generated_date": datetime.date.today().strftime("%Y年%m月%d日"),
    "benchmark_source": "Baymard Institute 2026 / CXL Institute",

    # -- KPI --
    "total_sessions": 0,
    "total_users": 0,
    "user_breakdown": "Desktop 0 / Mobile 0 / Tablet 0",
    "mobile_ratio": "0%",
    "overall_bounce_rate": "0%",
    "bounce_rate_eval": "",
    "total_key_events": 0,
    "overall_cvr": "0%",
    "cvr_eval": "",
    "estimated_monthly_loss": "約0リード",
    "estimated_recovery": "0件回復見込み",

    # -- Graph 1: 損失リードマップ --
    "lost_leads": {
        "bottlenecks": [
            # (label, lost_now, recovered, calc_note)
        ],
    },

    # -- Graph 2: ファネルウォーターフォール --
    "funnel": {
        "steps": ["page_view", "scroll", "cta_click", "demo_start",
                  "demo_\ncomplete", "generate_\nlead", "form_start", "form_submit"],
        "desktop": [],
        "mobile": [],
        "mobile_critical_annotation": {
            "text": "",
            "xy": (5.5, 1),
            "xytext": (6.5, 10),
        },
    },

    # -- Graph 3: CTA導線ヒートマップ --
    "cta_matrix": {
        "pages": [],       # list of page paths
        "pvs": [],         # list of PV counts
        "events": ["cta_click", "demo_start", "demo_\ncomplete",
                   "generate_\nlead", "form_start", "form_submit"],
        "matrix": [],      # 2D list: pages x events
    },

    # -- Graph 4: フォーム離脱ファネル --
    "form_funnel": {
        "steps": [],    # ["ページ閲覧", "generate_lead", "form_start", "form_submit"]
        "users": [],    # [36, 20, 2, 2]
        "actions": [],  # drop description for each step
        "fixes": [],    # improvement action for each step
    },

    # -- Graph 5: チャネル品質マトリクス --
    "channels": {
        # (name, sessions, avg_duration_sec, cvr_pct, color)
        "sources": [],
        "median_duration": 250,
        "median_cvr": 5,
        "callouts": [],  # list of (xy, xytext, text, color)
    },

    # -- Graph 6: ICE優先度 --
    "proposals": {
        # Each: (name, summary, impact, confidence, ease)
        "items": [],
    },

    # -- Graph 7: デバイス×ページ直帰率ギャップ --
    "bounce_gap": {
        "pages": [],
        "desktop": [],
        "mobile": [],  # None for missing data
        "benchmark": 40,
    },
}

# ================================================================
# NARRATIVES INJECTION POINT
# Phase 4-6 の分析文章をここに注入する
# ================================================================
NARRATIVES = {
    "executive_summary": [
        # list of paragraph strings
    ],
    "executive_accent": "",  # one-line accent text (e.g. estimated loss)

    "lost_leads_intro": "",
    "lost_leads_calc_intro": "",
    "lost_leads_conclusion": "",

    "funnel_intro": "",
    "funnel_desktop": [],      # list of paragraphs
    "funnel_mobile": [],       # list of paragraphs
    "funnel_mobile_accent": "",

    "cta_intro": "",
    "cta_analysis": [],        # list of paragraphs
    "cta_table": {
        "headers": ["ページ", "月間PV", "追加するCTA", "設置位置", "文言"],
        "rows": [],
        "col_widths": [2.5, 1.5, 3, 4, 5],
    },
    "cta_note": "",

    "form_intro": "",
    "form_problem": [],       # list of paragraphs
    "form_current_table": {
        "headers": ["フィールド", "必須/任意", "入力タイプ"],
        "rows": [],
        "col_widths": [4, 3, 4],
    },
    "form_current_note": "",
    "form_improvements": [],   # list of (bold_title, paragraph)
    "form_before_after": {
        "headers": ["項目", "Before", "After"],
        "rows": [],
        "col_widths": [3.5, 5.5, 5.5],
    },

    "channel_intro": "",
    "channel_sections": [],    # list of {"title": str, "paragraphs": [], "accent": str|None}
    "channel_filter_note": "",

    "bounce_intro": "",
    "bounce_analysis": [],

    "ice_intro": "",
    "ice_table": {
        "headers": ["#", "施策", "I", "C", "E", "ICE", "着手"],
        "rows": [],
        "col_widths": [0.8, 7, 1, 1, 1, 1, 2],
    },
    "ice_rationale_title": "",
    "ice_rationale": [],

    "roadmap_intro": "",
    "roadmap_table": {
        "headers": ["週", "施策", "ICE", "作業内容", "前提条件"],
        "rows": [],
        "col_widths": [1.8, 4, 1, 5, 3],
    },
    "roadmap_note": "",
}

# -- Output config --
OUTPUT_DIR = os.getcwd()
IMG_DIR = "/tmp/cro_charts"


# ================================================================
# CHART GENERATION FUNCTIONS
# ================================================================

def _save_chart(fig, name):
    os.makedirs(IMG_DIR, exist_ok=True)
    path = os.path.join(IMG_DIR, f"{name}.png")
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white", edgecolor="none")
    plt.close(fig)
    return path


def chart_lost_leads():
    """Graph 1: 損失リードマップ（横棒グラフ）"""
    d = DATA["lost_leads"]
    if not d["bottlenecks"]:
        return None

    labels = [b[0] for b in d["bottlenecks"]]
    lost_now = [b[1] for b in d["bottlenecks"]]
    recovered = [b[2] for b in d["bottlenecks"]]
    calc_notes = [b[3] for b in d["bottlenecks"]]

    fig, ax = plt.subplots(figsize=(14, max(5, len(labels) * 1.1)))
    y_pos = np.arange(len(labels))

    ax.barh(y_pos, lost_now, height=0.55, color=C["red"], alpha=0.85,
            edgecolor="white", linewidth=0.5, label="現在の月間損失リード数")
    ax.barh(y_pos, recovered, height=0.55,
            left=[l * 1.05 for l in lost_now],
            color=C["green"], alpha=0.7, edgecolor="white", linewidth=0.5,
            label="改善後の月間回復見込み")

    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=10, linespacing=1.3)
    ax.invert_yaxis()
    ax.set_xlabel("リード数/月", fontsize=12, fontweight="bold")
    ax.set_title("損失リードマップ — どこで何件のリードが消えているか",
                 fontsize=15, fontweight="bold", pad=15)

    for i, (lost, recov, note) in enumerate(zip(lost_now, recovered, calc_notes)):
        ax.text(lost/2, i, f"-{lost:.0f}件/月", ha="center", va="center",
                fontsize=11, fontweight="bold", color="white")
        ax.text(lost * 1.05 + recov/2, i, f"+{recov:.0f}件", ha="center", va="center",
                fontsize=10, fontweight="bold", color="white")
        ax.text(lost * 1.05 + recov + 0.5, i, note,
                fontsize=8, va="center", color=C["dark"], linespacing=1.4)

    max_x = max(l * 1.05 + r + 15 for l, r in zip(lost_now, recovered))
    ax.set_xlim(0, max_x)
    ax.legend(loc="lower right", fontsize=10, framealpha=0.9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    total_lost = sum(lost_now)
    total_recov = sum(recovered)
    ax.text(0.98, 0.02,
            f"合計損失: 月{total_lost:.0f}件 → 改善後: +{total_recov:.0f}件回復見込み",
            transform=ax.transAxes, ha="right", va="bottom",
            fontsize=12, fontweight="bold", color=C["accent"],
            bbox=dict(boxstyle="round,pad=0.4", facecolor=C["red_bg"], edgecolor=C["accent"]))

    fig.tight_layout(pad=2)
    return _save_chart(fig, "01_lost_leads")


def chart_funnel_waterfall():
    """Graph 2: ファネルウォーターフォール（Desktop / Mobile 2段）"""
    d = DATA["funnel"]
    if not d["desktop"] or not d["mobile"]:
        return None

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(14, 10), sharex=True)
    steps = d["steps"]
    desktop = d["desktop"]
    mobile = d["mobile"]

    def draw_waterfall(ax, data, label, color_main):
        x = np.arange(len(steps))
        bars = ax.bar(x, data, color=color_main, width=0.5, edgecolor="white", linewidth=0.5)
        for i, bar in enumerate(bars):
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width()/2, h + max(data)*0.02,
                        str(int(h)), ha="center", fontsize=10, fontweight="bold", color=color_main)
            if i < len(data) - 1 and data[i] > 0:
                drop = data[i] - data[i+1]
                drop_pct = drop / data[i] * 100
                if drop_pct > 30:
                    severity = "[!!]" if drop_pct > 80 else "[!]"
                    txt_color = C["red"] if drop_pct > 50 else C["yellow"]
                    ax.annotate(
                        f"{severity} -{drop}人\n({drop_pct:.0f}%離脱)",
                        xy=(i + 0.5, (data[i] + data[i+1]) / 2),
                        fontsize=8, ha="center", va="center",
                        color=txt_color, fontweight="bold",
                        bbox=dict(boxstyle="round,pad=0.2",
                                  facecolor=C["red_bg"] if drop_pct > 50 else C["yellow_bg"],
                                  edgecolor=txt_color, alpha=0.9))
        ax.set_xticks(x)
        ax.set_xticklabels(steps, fontsize=9)
        ax.set_ylabel("イベント数", fontsize=11)
        ax.set_title(f"{label} ファネル", fontsize=13, fontweight="bold", pad=10)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.set_yscale("symlog", linthresh=1)
        ax.set_ylim(0, max(data) * 1.3)

    draw_waterfall(ax1, desktop, "Desktop", C["primary"])
    draw_waterfall(ax2, mobile, "Mobile", C["accent"])

    crit = d.get("mobile_critical_annotation")
    if crit and crit["text"]:
        ax2.annotate(
            crit["text"],
            xy=crit["xy"], xytext=crit["xytext"],
            arrowprops=dict(arrowstyle="->", color=C["red"], lw=2.5),
            fontsize=11, fontweight="bold", color=C["red"],
            bbox=dict(boxstyle="round,pad=0.4", facecolor=C["red_bg"], edgecolor=C["red"]))

    fig.tight_layout(pad=2)
    return _save_chart(fig, "02_funnel_waterfall")


def chart_cta_heatmap():
    """Graph 3: CTA導線ヒートマップ"""
    d = DATA["cta_matrix"]
    if not d["pages"]:
        return None

    pages = d["pages"]
    pvs = d["pvs"]
    events = d["events"]
    matrix = np.array(d["matrix"], dtype=float)

    fig, ax = plt.subplots(figsize=(14, max(4, len(pages) * 0.8)))

    for i in range(len(pages)):
        for j in range(len(events)):
            val = matrix[i, j]
            if val == 0:
                color = C["red_bg"]
                txt, txt_color = "0", C["red"]
            else:
                intensity = min(val / 30, 1.0)
                color = C["green_bg"] if intensity > 0.3 else C["yellow_bg"]
                txt = str(int(val))
                txt_color = C["green"] if intensity > 0.3 else C["dark"]
            rect = plt.Rectangle((j, i), 1, 1, facecolor=color, edgecolor="white", linewidth=2)
            ax.add_patch(rect)
            ax.text(j + 0.5, i + 0.5, txt, ha="center", va="center",
                    fontsize=12, fontweight="bold", color=txt_color)

    pv_x = len(events) + 0.3
    max_pv = max(pvs) if pvs else 1
    for i, (page, pv) in enumerate(zip(pages, pvs)):
        bar_width = (pv / max_pv) * 3
        rect = plt.Rectangle((pv_x, i + 0.15), bar_width, 0.7,
                              facecolor=C["primary"], alpha=0.7, edgecolor="none")
        ax.add_patch(rect)
        ax.text(pv_x + bar_width + 0.1, i + 0.5, f"{pv} PV",
                fontsize=9, va="center", color=C["dark"])

    ax.set_xlim(-0.5, pv_x + 4.5)
    ax.set_ylim(-0.5, len(pages) + 0.5)
    ax.set_xticks([j + 0.5 for j in range(len(events))] + [pv_x + 1.5])
    ax.set_xticklabels(events + ["PV"], fontsize=10)
    ax.set_yticks([i + 0.5 for i in range(len(pages))])
    ax.set_yticklabels(pages, fontsize=11, fontweight="bold")
    ax.invert_yaxis()
    ax.set_title("CTA導線ヒートマップ — どのページにどのCV導線があるか",
                 fontsize=14, fontweight="bold", pad=12)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(left=False, bottom=False)

    legend_elements = [
        mpatches.Patch(facecolor=C["red_bg"], edgecolor=C["red"], label="0件 = CTA死角 (即追加)"),
        mpatches.Patch(facecolor=C["yellow_bg"], edgecolor=C["yellow"], label="1-5件 = 改善余地"),
        mpatches.Patch(facecolor=C["green_bg"], edgecolor=C["green"], label="6件+ = 機能中"),
    ]
    ax.legend(handles=legend_elements, loc="upper right", fontsize=9, framealpha=0.9)
    fig.tight_layout()
    return _save_chart(fig, "03_cta_heatmap")


def chart_form_funnel():
    """Graph 4: フォーム離脱ファネル（台形）"""
    d = DATA["form_funnel"]
    if not d["steps"]:
        return None

    steps = d["steps"]
    users = d["users"]
    actions = d["actions"]
    fixes = d["fixes"]

    fig, ax = plt.subplots(figsize=(12, 7))
    max_w = 10
    center_x = 6
    total_h = 12
    step_h = total_h / len(steps)

    for i, (step, user, action, fix) in enumerate(zip(steps, users, actions, fixes)):
        w = max_w * (user / users[0]) if users[0] > 0 else 0
        y = total_h - (i + 1) * step_h
        w_next = max_w * (users[i+1] / users[0]) if i < len(steps) - 1 and users[0] > 0 else w

        verts = [
            (center_x - w/2, y + step_h),
            (center_x + w/2, y + step_h),
            (center_x + w_next/2, y),
            (center_x - w_next/2, y),
        ]
        color = C["green"] if i == 0 else (C["yellow"] if i == 1 else (C["red"] if user <= 2 and i >= 2 else C["green"]))
        poly = Polygon(verts, closed=True, facecolor=color, alpha=0.2, edgecolor=color, linewidth=2)
        ax.add_patch(poly)

        ax.text(0.5, y + step_h/2, f"{step}\n{user}人", fontsize=12, fontweight="bold",
                va="center", ha="left", color=C["dark"])

        if action:
            bg = C["red_bg"] if "90%" in action else (C["yellow_bg"] if "44%" in action else C["green_bg"])
            txt_c = C["red"] if "90%" in action else (C["yellow"] if "44%" in action else C["green"])
            ax.text(center_x, y + step_h * 0.3,
                    action, fontsize=10, fontweight="bold",
                    va="center", ha="center", color=txt_c,
                    bbox=dict(boxstyle="round,pad=0.3", facecolor=bg, edgecolor=txt_c, alpha=0.9))

        if fix and i > 0:
            ax.text(12.5, y + step_h/2, fix, fontsize=9, va="center", ha="left",
                    color=C["dark"], linespacing=1.4,
                    bbox=dict(boxstyle="round,pad=0.3", facecolor="#f0f0f0", edgecolor="#cccccc"))

    ax.set_xlim(-0.5, 22)
    ax.set_ylim(-1, total_h + 1)
    ax.set_title("/contact フォーム離脱ファネル + 改善策", fontsize=14, fontweight="bold", pad=12)
    ax.axis("off")
    fig.tight_layout()
    return _save_chart(fig, "04_form_funnel")


def chart_channel_quality():
    """Graph 5: チャネル品質マトリクス（4象限バブルチャート）"""
    d = DATA["channels"]
    if not d["sources"]:
        return None

    fig, ax = plt.subplots(figsize=(14, 8))
    median_dur = d["median_duration"]
    median_cvr = d["median_cvr"]

    ax.axvline(x=median_dur, color=C["gray"], linestyle="--", alpha=0.4)
    ax.axhline(y=median_cvr, color=C["gray"], linestyle="--", alpha=0.4)

    ax.text(median_dur * 0.3, 28, "効率的\n(ウォームトラフィック)", fontsize=10,
            color=C["primary"], alpha=0.5, ha="center", style="italic")
    ax.text(median_dur * 2.2, 28, "理想的\n(拡大すべき)", fontsize=10,
            color=C["green"], alpha=0.5, ha="center", style="italic")
    ax.text(median_dur * 0.3, 1.5, "要改善\n(LP/導線見直し)", fontsize=10,
            color=C["red"], alpha=0.5, ha="center", style="italic")
    ax.text(median_dur * 2.2, 1.5, "CTA不足\n(エンゲージメントは高い)", fontsize=10,
            color=C["yellow"], alpha=0.5, ha="center", style="italic")

    for name, sessions, dur, cvr, color in d["sources"]:
        ax.scatter(dur, cvr, s=sessions * 4, c=color, alpha=0.7,
                   edgecolors="white", linewidths=1.5, zorder=5)
        offset_y = 1.5 if cvr > median_cvr else -1.5
        ax.annotate(f"{name}\n({sessions}セッション)",
                    (dur, cvr), textcoords="offset points",
                    xytext=(10, offset_y * 5), fontsize=9, color=C["dark"],
                    arrowprops=dict(arrowstyle="-", color=C["gray"], alpha=0.3))

    for callout in d.get("callouts", []):
        xy, xytext, text, col = callout
        ax.annotate(text, xy=xy, xytext=xytext,
                    arrowprops=dict(arrowstyle="->", color=col, lw=2),
                    fontsize=10, fontweight="bold", color=col,
                    bbox=dict(boxstyle="round,pad=0.3",
                              facecolor=C.get(f"{col}_bg", "#f0f0f0") if col in C else "#f0f0f0",
                              edgecolor=col))

    ax.set_xlabel("平均滞在時間 (秒)", fontsize=12, fontweight="bold")
    ax.set_ylabel("CVR (%)", fontsize=12, fontweight="bold")
    ax.set_title(f"流入元の質マトリクス — どのチャネルに投資すべきか\n(バブルサイズ = セッション数)",
                 fontsize=14, fontweight="bold", pad=15)
    max_dur = max(s[2] for s in d["sources"]) * 1.15
    max_cvr = max(s[3] for s in d["sources"]) * 1.3
    ax.set_xlim(0, max(max_dur, median_dur * 3))
    ax.set_ylim(-2, max(max_cvr, 35))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _save_chart(fig, "05_channel_quality")


def chart_ice_priority():
    """Graph 6: ICE優先度（積み上げ横棒）"""
    d = DATA["proposals"]
    if not d["items"]:
        return None

    items = d["items"]
    proposals = [x[0] for x in items]
    summaries = [x[1] for x in items]
    impact = [x[2] for x in items]
    confidence = [x[3] for x in items]
    ease = [x[4] for x in items]
    ice = [(i+c+e)/3 for i, c, e in zip(impact, confidence, ease)]

    fig, ax = plt.subplots(figsize=(14, max(5, len(proposals) * 1.0)))
    y_pos = np.arange(len(proposals))

    ax.barh(y_pos, impact, height=0.6, color=C["accent"], alpha=0.85,
            label="Impact", edgecolor="white")
    ax.barh(y_pos, confidence, height=0.6, left=impact,
            color=C["primary"], alpha=0.85, label="Confidence", edgecolor="white")
    ax.barh(y_pos, ease, height=0.6,
            left=[i+c for i, c in zip(impact, confidence)],
            color=C["green"], alpha=0.85, label="Ease", edgecolor="white")

    ax.set_yticks(y_pos)
    ax.set_yticklabels(proposals, fontsize=11, fontweight="bold")
    ax.invert_yaxis()

    for i, (score, summary) in enumerate(zip(ice, summaries)):
        total = impact[i] + confidence[i] + ease[i]
        if score >= 8:
            bg, fg, timeline = C["green_bg"], C["green"], "今すぐ"
        elif score >= 7:
            bg, fg, timeline = C["yellow_bg"], C["yellow"], "今月中"
        else:
            bg, fg, timeline = "#f0f0f0", C["gray"], "来月以降"
        ax.text(total + 0.5, i - 0.15, f"ICE {score:.1f}",
                fontsize=12, fontweight="bold", color=fg, va="center",
                bbox=dict(boxstyle="round,pad=0.2", facecolor=bg, edgecolor=fg))
        ax.text(total + 0.5, i + 0.2, f"[{timeline}] {summary}",
                fontsize=8, color=C["dark"], va="center")

    ax.set_xlabel("スコア合計 (Impact + Confidence + Ease)", fontsize=11)
    ax.set_xlim(0, 35)
    ax.set_title("ICE優先度 — 何から着手するか", fontsize=14, fontweight="bold", pad=12)
    ax.legend(loc="lower right", fontsize=10, framealpha=0.9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _save_chart(fig, "06_ice_priority")


def chart_bounce_gap():
    """Graph 7: デバイス×ページ直帰率ギャップ"""
    d = DATA["bounce_gap"]
    if not d["pages"]:
        return None

    pages = d["pages"]
    desktop_b = d["desktop"]
    mobile_b = d["mobile"]
    benchmark = d["benchmark"]

    fig, ax = plt.subplots(figsize=(12, max(4, len(pages) * 1.2)))
    y_pos = np.arange(len(pages))
    height = 0.3

    ax.barh(y_pos - height/2, desktop_b, height, label="Desktop",
            color=C["primary"], edgecolor="white")
    mobile_vals = [v if v is not None else 0 for v in mobile_b]
    bars_m = ax.barh(y_pos + height/2, mobile_vals, height, label="Mobile",
                     color=C["accent"], edgecolor="white")

    for i, v in enumerate(mobile_b):
        if v is None:
            bars_m[i].set_color(C["light"])
            bars_m[i].set_edgecolor(C["gray"])

    ax.axvline(x=benchmark, color=C["red"], linestyle="--", alpha=0.6, linewidth=1.5)
    ax.text(benchmark + 1, -0.5, f"要改善ライン ({benchmark}%)", fontsize=9, color=C["red"])

    for i, (d_val, m_val) in enumerate(zip(desktop_b, mobile_b)):
        ax.text(d_val + 1, i - height/2, f"{d_val:.1f}%", va="center", fontsize=10,
                fontweight="bold", color=C["primary"])
        if m_val is not None:
            ax.text(m_val + 1, i + height/2, f"{m_val:.1f}%", va="center", fontsize=10,
                    fontweight="bold", color=C["accent"])
            gap = m_val - d_val
            if gap >= 15:
                ax.text(max(d_val, m_val) + 12, i, f"Gap {gap:.1f}%",
                        fontsize=10, fontweight="bold", color=C["red"], va="center",
                        bbox=dict(boxstyle="round,pad=0.2", facecolor=C["red_bg"], edgecolor=C["red"]))
        else:
            ax.text(3, i + height/2, "データ不足", va="center", fontsize=8, color=C["gray"])

        if d_val >= 100:
            ax.text(d_val - 5, i - height/2, "即離脱", va="center", ha="right",
                    fontsize=9, fontweight="bold", color="white",
                    bbox=dict(boxstyle="round,pad=0.2", facecolor=C["red"]))
        elif d_val >= 50:
            ax.text(d_val - 5, i - height/2, "危険", va="center", ha="right",
                    fontsize=9, fontweight="bold", color="white",
                    bbox=dict(boxstyle="round,pad=0.2", facecolor=C["yellow"]))

    ax.set_yticks(y_pos)
    ax.set_yticklabels(pages, fontsize=12, fontweight="bold")
    ax.invert_yaxis()
    ax.set_xlabel("直帰率 (%)", fontsize=11)
    ax.set_xlim(0, 120)
    ax.set_title("デバイス×LP直帰率ギャップ（15%以上 = モバイル最適化必須）",
                 fontsize=14, fontweight="bold", pad=12)
    ax.legend(fontsize=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _save_chart(fig, "07_bounce_gap")


# ================================================================
# WORD DOCUMENT GENERATION
# ================================================================

def generate_docx(chart_paths):
    """Generate Word report with charts and narratives."""
    doc = Document()

    # -- Style setup --
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Hiragino Sans"
        hs.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
        if level == 1:
            hs.font.size = Pt(20)
            hs.paragraph_format.space_before = Pt(24)
        elif level == 2:
            hs.font.size = Pt(15)
            hs.paragraph_format.space_before = Pt(18)
        else:
            hs.font.size = Pt(12)
            hs.paragraph_format.space_before = Pt(12)

    # -- Helpers --
    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            cell = t.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.rows[i+1].cells[j]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        if col_widths:
            for j, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[j].width = Cm(w)
        doc.add_paragraph("")

    def add_img(key, width=Inches(6.2)):
        path = chart_paths.get(key)
        if path and os.path.exists(path):
            doc.add_picture(path, width=width)
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    def add_bold_para(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        return p

    def add_accent_para(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)
        run.font.size = Pt(11)
        return p

    N = NARRATIVES

    # ── Title page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DATA["site_name"])
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CROディープ分析レポート")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GA4クリーンデータ × 実ページUX構造の突き合わせ分析")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    meta_items = [
        f"分析期間: {DATA['period']}",
        f"対象サイト: {DATA['site_url']}",
        f"GA4プロパティ: {DATA['ga4_property']}",
        f"生成日: {DATA['generated_date']}",
        f"ベンチマーク出典: {DATA['benchmark_source']}",
    ]
    for item in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("エグゼクティブサマリー", level=1)
    for para in N["executive_summary"]:
        doc.add_paragraph(para)
    if N["executive_accent"]:
        add_accent_para(N["executive_accent"])

    add_table(
        ["指標", "数値", "評価"],
        [
            ["総セッション", str(DATA["total_sessions"]), "—"],
            ["総ユーザー", f"{DATA['total_users']} ({DATA['user_breakdown']})", f"モバイル比率{DATA['mobile_ratio']}"],
            ["全体直帰率", DATA["overall_bounce_rate"], DATA["bounce_rate_eval"]],
            ["全体CVR", f"{DATA['overall_cvr']} ({DATA['total_key_events']} keyEvents)", DATA["cvr_eval"]],
            ["月間推定損失", DATA["estimated_monthly_loss"], DATA["estimated_recovery"]],
        ],
        col_widths=[4, 8, 6],
    )
    doc.add_page_break()

    # ── Section 1: 損失リードマップ ──
    doc.add_heading("1. 損失リードマップ — どこで何件のリードが消えているか", level=1)
    if N["lost_leads_intro"]:
        doc.add_paragraph(N["lost_leads_intro"])
    add_img("lost_leads")
    if N["lost_leads_calc_intro"]:
        doc.add_heading("算出根拠", level=3)
        doc.add_paragraph(N["lost_leads_calc_intro"])
    ll = DATA["lost_leads"]["bottlenecks"]
    if ll:
        add_table(
            ["ボトルネック", "損失/月", "回復見込み", "算出根拠"],
            [[b[0].replace("\n", " "), f"{b[1]:.0f}件", f"+{b[2]:.0f}件", b[3].replace("\n", " ")] for b in ll],
            col_widths=[5, 2, 2, 8],
        )
    if N["lost_leads_conclusion"]:
        doc.add_paragraph(N["lost_leads_conclusion"])
    doc.add_page_break()

    # ── Section 2: ファネル分析 ──
    doc.add_heading("2. ファネル分析 — Desktop と Mobile で何が違うか", level=1)
    if N["funnel_intro"]:
        doc.add_paragraph(N["funnel_intro"])
    add_img("funnel")

    if N["funnel_desktop"]:
        doc.add_heading("Desktop のファネル", level=2)
        for para in N["funnel_desktop"]:
            doc.add_paragraph(para)
    if N["funnel_mobile"]:
        doc.add_heading("Mobile のファネル — 致命的な断裂", level=2)
        for para in N["funnel_mobile"]:
            doc.add_paragraph(para)
    if N["funnel_mobile_accent"]:
        add_accent_para(N["funnel_mobile_accent"])

    add_table(
        ["離脱率", "深刻度", "対応"],
        [["> 80%", "致命的", "即時対応。原因はページ構造にある"],
         ["50-80%", "要改善", "A/Bテストで検証後に改修"],
         ["30-50%", "許容範囲", "改善余地あり、優先度は中"],
         ["< 30%", "正常", "他の施策を優先"]],
        col_widths=[3, 3, 8],
    )
    doc.add_page_break()

    # ── Section 3: CTA導線ヒートマップ ──
    doc.add_heading("3. CTA導線の死角 — どのページに何が足りないか", level=1)
    if N["cta_intro"]:
        doc.add_paragraph(N["cta_intro"])
    add_img("cta_heatmap")
    for para in N.get("cta_analysis", []):
        doc.add_paragraph(para)
    ct = N.get("cta_table", {})
    if ct.get("rows"):
        doc.add_heading("各ページへのCTA追加仕様", level=2)
        add_table(ct["headers"], ct["rows"], ct.get("col_widths"))
    if N.get("cta_note"):
        doc.add_paragraph(N["cta_note"])
    doc.add_page_break()

    # ── Section 4: フォーム改善 ──
    doc.add_heading("4. /contact フォーム改善 — 最大のボトルネック", level=1)
    if N["form_intro"]:
        doc.add_paragraph(N["form_intro"])
    add_img("form_funnel")
    if N.get("form_problem"):
        doc.add_heading("問題の構造", level=2)
        for para in N["form_problem"]:
            doc.add_paragraph(para)
    fc = N.get("form_current_table", {})
    if fc.get("rows"):
        doc.add_heading("現在のフォーム構造（WebFetchで確認）", level=2)
        add_table(fc["headers"], fc["rows"], fc.get("col_widths"))
    if N.get("form_current_note"):
        doc.add_paragraph(N["form_current_note"])
    if N.get("form_improvements"):
        doc.add_heading("改善仕様", level=2)
        for title, body in N["form_improvements"]:
            add_bold_para(title)
            doc.add_paragraph(body)
    fba = N.get("form_before_after", {})
    if fba.get("rows"):
        add_table(fba["headers"], fba["rows"], fba.get("col_widths"))
    doc.add_page_break()

    # ── Section 5: チャネル戦略 ──
    doc.add_heading("5. チャネル品質分析 — どの流入元に投資すべきか", level=1)
    if N["channel_intro"]:
        doc.add_paragraph(N["channel_intro"])
    add_img("channel")
    for section in N.get("channel_sections", []):
        doc.add_heading(section["title"], level=2)
        for para in section.get("paragraphs", []):
            doc.add_paragraph(para)
        if section.get("accent"):
            add_accent_para(section["accent"])
    if N.get("channel_filter_note"):
        doc.add_heading("フィルタ修正事項", level=3)
        doc.add_paragraph(N["channel_filter_note"])
    doc.add_page_break()

    # ── Section 6: 直帰率ギャップ ──
    doc.add_heading("6. デバイス別直帰率ギャップ", level=1)
    add_img("bounce")
    if N.get("bounce_intro"):
        doc.add_paragraph(N["bounce_intro"])
    for para in N.get("bounce_analysis", []):
        doc.add_paragraph(para)
    doc.add_page_break()

    # ── Section 7: ICE優先度 ──
    doc.add_heading("7. 改善施策の優先順位（ICEスコア）", level=1)
    if N["ice_intro"]:
        doc.add_paragraph(N["ice_intro"])
    add_img("ice")
    it = N.get("ice_table", {})
    if it.get("rows"):
        add_table(it["headers"], it["rows"], it.get("col_widths"))
    if N.get("ice_rationale_title"):
        doc.add_heading(N["ice_rationale_title"], level=2)
    for para in N.get("ice_rationale", []):
        doc.add_paragraph(para)
    doc.add_page_break()

    # ── Section 8: ロードマップ ──
    doc.add_heading("8. 着手ロードマップ", level=1)
    if N["roadmap_intro"]:
        doc.add_paragraph(N["roadmap_intro"])
    rt = N.get("roadmap_table", {})
    if rt.get("rows"):
        add_table(rt["headers"], rt["rows"], rt.get("col_widths"))
    if N.get("roadmap_note"):
        doc.add_paragraph(N["roadmap_note"])

    # ── Save ──
    output_name = f"{DATA['site_name']}_CRO分析レポート.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    return output_path


# ================================================================
# MAIN
# ================================================================

def main():
    print("Generating charts...")
    chart_paths = {}
    chart_funcs = [
        ("lost_leads", chart_lost_leads),
        ("funnel", chart_funnel_waterfall),
        ("cta_heatmap", chart_cta_heatmap),
        ("form_funnel", chart_form_funnel),
        ("channel", chart_channel_quality),
        ("ice", chart_ice_priority),
        ("bounce", chart_bounce_gap),
    ]
    for name, fn in chart_funcs:
        print(f"  {name}...")
        path = fn()
        if path:
            chart_paths[name] = path

    print(f"Generated {len(chart_paths)} charts")

    print("Generating Word document...")
    output = generate_docx(chart_paths)
    print(f"Saved: {output}")
    return output


if __name__ == "__main__":
    main()
